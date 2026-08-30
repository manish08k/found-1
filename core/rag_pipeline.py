"""
RAG Document Ingestion Pipeline.

Handles: parsing (PDF, DOCX, TXT, Markdown, HTML), chunking (sliding window),
embedding (via provider config), and storage (VectorDocument rows).
"""
import io
import re
import uuid
from typing import Optional

import structlog
from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from storage.models import VectorDocument, DocumentStore

log = structlog.get_logger(__name__)


# ─── Parsing ──────────────────────────────────────────────────────────────────

def _parse_text(file_bytes: bytes, filename: str, file_type: str) -> str:
    """Extract plain text from supported file types."""
    ft = file_type.lower() if file_type else ""
    name_lower = filename.lower() if filename else ""

    if ft in ("text/plain",) or name_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")

    if ft in ("text/markdown",) or name_lower.endswith((".md", ".markdown")):
        return file_bytes.decode("utf-8", errors="replace")

    if ft in ("text/html", "application/xhtml+xml") or name_lower.endswith((".html", ".htm")):
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_bytes, "html.parser")
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # Fallback: strip tags with regex
            text = file_bytes.decode("utf-8", errors="replace")
            return re.sub(r"<[^>]+>", " ", text)

    if ft in ("application/pdf",) or name_lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            raise ValueError("PDF parsing requires the 'pypdf' package. Install with: pip install pypdf")

    if ft in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or name_lower.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ValueError("DOCX parsing requires 'python-docx'. Install with: pip install python-docx")

    # Fallback: try decoding as text
    return file_bytes.decode("utf-8", errors="replace")


# ─── Chunking ─────────────────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks using a sliding window."""
    if not text or not text.strip():
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]

        # Try to break at a sentence boundary near the end
        if end < text_len:
            last_period = chunk.rfind(". ")
            last_newline = chunk.rfind("\n")
            break_point = max(last_period, last_newline)
            if break_point > chunk_size * 0.5:
                chunk = chunk[: break_point + 1]
                end = start + break_point + 1

        if chunk.strip():
            chunks.append(chunk.strip())

        start = end - overlap
        if start <= 0 and end >= text_len:
            break
        if end >= text_len:
            break

    return chunks


# ─── Embedding ────────────────────────────────────────────────────────────────

async def _embed_texts(texts: list[str], model: str = "text-embedding-3-small") -> list[list[float]]:
    """Embed texts using OpenAI embeddings API."""
    from integrations.vector.handler import _embed
    return await _embed(texts, model)


# ─── Pipeline ─────────────────────────────────────────────────────────────────

async def ingest_document(
    store_id: str,
    file_bytes: bytes,
    filename: str,
    file_type: str,
    db: AsyncSession,
    settings: Optional[dict] = None,
) -> dict:
    """
    Full ingestion pipeline:
    1. Parse file to text
    2. Chunk with sliding window
    3. Embed each chunk
    4. Store as VectorDocument rows
    """
    # Load store config
    result = await db.execute(select(DocumentStore).where(DocumentStore.id == store_id))
    store = result.scalar_one_or_none()
    if not store:
        raise ValueError(f"Document store {store_id} not found")

    chunk_size = store.chunk_size or 1000
    chunk_overlap = store.chunk_overlap or 200
    embedding_model = store.embedding_model or "text-embedding-3-small"

    # 1. Parse
    text = _parse_text(file_bytes, filename, file_type)
    if not text.strip():
        return {"status": "empty", "chunks": 0, "filename": filename}

    # 2. Chunk
    chunks = _chunk_text(text, chunk_size, chunk_overlap)
    if not chunks:
        return {"status": "no_chunks", "chunks": 0, "filename": filename}

    # 3. Embed in batches of 100
    source_doc_id = str(uuid.uuid4())
    all_embeddings: list[list[float]] = []
    batch_size = 100
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        embeddings = await _embed_texts(batch, embedding_model)
        all_embeddings.extend(embeddings)

    # 4. Store
    doc_ids = []
    for idx, (chunk_text, embedding) in enumerate(zip(chunks, all_embeddings)):
        doc = VectorDocument(
            workflow_id=store_id,  # reuse workflow_id to scope to store
            collection=store_id,
            text=chunk_text,
            embedding=embedding,
            doc_metadata={"filename": filename, "file_type": file_type},
            source_document_id=source_doc_id,
            chunk_index=idx,
            metadata={"filename": filename, "chunk_index": idx, "total_chunks": len(chunks)},
        )
        db.add(doc)
        await db.flush()
        doc_ids.append(doc.id)

    await db.commit()

    return {
        "status": "success",
        "source_document_id": source_doc_id,
        "filename": filename,
        "chunks": len(chunks),
        "doc_ids": doc_ids,
    }


async def query_documents(
    store_id: str,
    query_text: str,
    top_k: int = 5,
    filters: Optional[dict] = None,
    db: AsyncSession = None,
    settings: Optional[dict] = None,
) -> dict:
    """
    Query pipeline:
    1. Embed query text
    2. Cosine similarity search across VectorDocument
    3. Apply metadata filters
    4. Return ranked results
    """
    from integrations.vector.handler import _embed, _cosine_similarity

    # Load store config
    result = await db.execute(select(DocumentStore).where(DocumentStore.id == store_id))
    store = result.scalar_one_or_none()
    if not store:
        raise ValueError(f"Document store {store_id} not found")

    embedding_model = store.embedding_model or "text-embedding-3-small"

    # 1. Embed query
    query_embedding = (await _embed([query_text], embedding_model))[0]

    # 2. Load all docs for this store
    stmt = select(VectorDocument).where(VectorDocument.collection == store_id)
    doc_result = await db.execute(stmt)
    docs = doc_result.scalars().all()

    # 3. Score and filter
    scored = []
    for doc in docs:
        # Apply metadata filters
        if filters:
            doc_meta = doc.doc_metadata or {}
            doc_extra = doc.metadata or {}
            merged = {**doc_meta, **doc_extra}
            skip = False
            for fk, fv in filters.items():
                if merged.get(fk) != fv:
                    skip = True
                    break
            if skip:
                continue

        score = _cosine_similarity(query_embedding, doc.embedding)
        scored.append((doc, score))

    scored.sort(key=lambda x: x[1], reverse=True)
    top = scored[:top_k]

    return {
        "results": [
            {
                "text": d.text,
                "score": round(s, 4),
                "metadata": d.doc_metadata,
                "source_document_id": d.source_document_id,
                "chunk_index": d.chunk_index,
                "id": d.id,
            }
            for d, s in top
        ],
        "context": "\n\n".join(d.text for d, _ in top),
        "total_searched": len(docs),
        "query": query_text,
    }
